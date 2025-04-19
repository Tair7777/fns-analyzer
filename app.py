import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import streamlit as st
import chardet
import re
import datetime
import io
import base64
import random
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from analyzer import LegalAnalyzer
from config import SCHEMA_DIR, SCHEMA_MAPPING
import xmlschema
from io import StringIO
import xml.dom.minidom

class DocumentDetector:
    def detect_document_type(self, content):
        if '<Файл ИдФайл="ON_NSCHFDOPPR' in content or '<СвСчФакт' in content:
            return 'upd'
        elif '<Файл ИдФайл="NO_NDFL' in content:
            return 'ndfl'
        elif '<ЕГРЮЛ' in content:
            return 'egroul'
        elif '<Invoice' in content or '<СчетФактура' in content:
            return 'invoice'
        else:
            return 'unknown'

class XSDValidator:
    def __init__(self, schema_path):
        self.schema = xmlschema.XMLSchema(schema_path)
        self.modifications = []
        self.last_error = ""

    def validate(self, xml_content):
        try:
            self.schema.validate(xml_content)
            return True, ""
        except Exception as e:
            self.last_error = str(e)
            return False, str(e)
    
    def generate_xml_sample(self):
        try:
            placeholders = {
                'string': lambda: 'Пример значения',
                'integer': lambda: str(random.randint(1, 1000)),
                'date': lambda: datetime.datetime.now().strftime('%Y-%m-%d'),
                'ИНН': lambda: ''.join(random.choices('0123456789', k=12)),
                'КПП': lambda: ''.join(random.choices('0123456789', k=9)),
            }
            global_elements = [
                elem for elem in self.schema.elements.values() 
                if isinstance(elem, xmlschema.XsdElement)
            ]
            if not global_elements:
                return "Схема не содержит корневых элементов"
            root_element = global_elements[0]
            xml_dict = {root_element.name: self._fill_element(root_element.type, placeholders)}
            xml_string = xmlschema.to_xml(xml_dict)
            dom = xml.dom.minidom.parseString(xml_string)
            return dom.toprettyxml(indent="  ")
        except Exception as e:
            return f"Ошибка генерации: {str(e)}"
    
    def _fill_element(self, element, placeholders, depth=0):
        if depth > 2: 
            return "..."
        result = {}
        if hasattr(element, 'attributes'):
            for attr_name, attr in element.attributes.items():
                result[f"@{attr_name}"] = placeholders.get(attr.type.name, placeholders['string'])()
        if hasattr(element, 'content'):
            content_model = element.content
            if isinstance(content_model, xmlschema.XsdGroup):
                for child in content_model.iter_components():
                    if isinstance(child, xmlschema.XsdElement):
                        result[child.name] = self._fill_element(child.type, placeholders, depth+1)
                    elif isinstance(child, xmlschema.XsdGroup):
                        result.update(self._fill_element(child, placeholders, depth+1))
        return result if result else placeholders['string']()

def read_uploaded_file(uploaded_file):
    try:
        raw = uploaded_file.getvalue()
        first_200 = raw[:200].decode(errors='ignore')
        match = re.search(r'encoding=["\']([\w-]+)["\']', first_200)
        encoding = match.group(1) if match else chardet.detect(raw)['encoding'] or 'utf-8'
        return raw.decode(encoding)
    except Exception as e:
        st.error(f"Ошибка декодирования: {str(e)}")
        return None

def create_word_doc(content, filename):
    doc = Document()
    doc.add_heading('Анализ документа', 0)
    doc.add_paragraph(content)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_pdf(content, filename):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.drawString(72, 750, "Анализ документа")
    text = p.beginText(72, 730)
    for line in content.split('\n'):
        text.textLine(line)
    p.drawText(text)
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.getvalue()

def display_chat(content):
    st.subheader("Диалог с документом")
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
    
    if prompt := st.chat_input("Ваш вопрос..."):
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.spinner("Генерация ответа..."):
            try:
                response = analyzer.answer_question(content[:3000], prompt)
                st.session_state.chat_history.append({"role": "assistant", "content": response})
                st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {str(e)}")

st.set_page_config(page_title="Анализатор ФНС", layout="wide")

if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'current_content' not in st.session_state:
    st.session_state.current_content = None

analyzer = LegalAnalyzer()

st.title("🔍 Анализатор документов ФНС")
uploaded_file = st.file_uploader("Загрузите XML/XSD файл", type=["xml", "xsd"])

if uploaded_file:
    content = read_uploaded_file(uploaded_file)
    if content:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        
        # Обработка XSD файлов
        if file_ext == 'xsd':
            try:
                validator = XSDValidator(StringIO(content))
                example_xml = validator.generate_xml_sample()
                
                st.subheader("Пример XML для этой схемы")
                st.code(example_xml, language='xml')
                
                st.download_button(
                    label="📥 Скачать пример XML",
                    data=example_xml,
                    file_name="example.xml",
                    mime="text/xml"
                )

                if st.button("🔍 Проанализировать этот пример"):
                    st.session_state.chat_history = []
                    st.session_state.current_content = example_xml
                    with st.spinner("Анализ..."):
                        try:
                            st.session_state.analysis_result = analyzer.analyze_text(example_xml)
                        except Exception as e:
                            st.error(f"Ошибка анализа: {str(e)}")

            except Exception as e:
                st.error(f"Ошибка обработки XSD: {str(e)}")
        
        # Обработка XML файлов
        else:
            st.session_state.current_content = content
            detector = DocumentDetector()
            doc_type = detector.detect_document_type(content)
            schema_path = os.path.join(SCHEMA_DIR, SCHEMA_MAPPING.get(doc_type, ''))
            if os.path.exists(schema_path):
                validator = XSDValidator(schema_path)
                is_valid, error = validator.validate(content)
                if not is_valid:
                    st.warning(f"Ошибки валидации: {error}")

            if st.button("Проанализировать документ"):
                st.session_state.chat_history = []
                with st.spinner("Анализ..."):
                    try:
                        st.session_state.analysis_result = analyzer.analyze_text(content)
                    except Exception as e:
                        st.error(f"Ошибка анализа: {str(e)}")

        # Отображение результатов и чата для любого типа
        if st.session_state.analysis_result:
            analysis_text = st.session_state.analysis_result.get('analysis', '') if isinstance(st.session_state.analysis_result, dict) else st.session_state.analysis_result
            
            col1, col2 = st.columns([3,1])
            with col1:
                st.subheader("Результаты анализа")
                st.markdown(analysis_text)
                doc_bytes = create_word_doc(analysis_text, uploaded_file.name)
                st.download_button(
                    label="📥 Скачать анализ (DOCX)",
                    data=doc_bytes,
                    file_name=f"analysis_{uploaded_file.name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col2:
                if st.session_state.current_content:
                    display_chat(st.session_state.current_content)
